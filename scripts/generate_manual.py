#!/usr/bin/env python
"""Génère le manuel de procédure utilisateur AvoLex (Word)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Manuel_Procedure_AvoLex.docx"

# Palette professionnelle
NAVY = RGBColor(0x1A, 0x2B, 0x4A)
ACCENT = RGBColor(0x2C, 0x5F, 0x8A)
GRAY = RGBColor(0x5A, 0x5A, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER = RGBColor(0xE8, 0xEE, 0xF4)


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run("1")
    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run(
        "Mettez à jour la table des matières : clic droit → Mettre à jour les champs."
    )
    run4.font.italic = True
    run4.font.color.rgb = GRAY

    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri Light"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY if level == 1 else ACCENT
        style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def setup_page(section, *, footer_text: str = "AvoLex — Manuel de procédure") -> None:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.orientation = WD_ORIENT.PORTRAIT

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = footer_text
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.font.italic = True

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)
    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY


def add_cover(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AvoLex")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Manuel de procédure utilisateur")
    run.font.name = "Calibri Light"
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Gestion de cabinet d'avocats — SaaS multi-tenant")
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Version 1.0 — {date.today():%d/%m/%Y}")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("Document interne — Usage cabinet")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY

    doc.add_page_break()


def add_info_box(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "E8EEF4")
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}\n")
    r.bold = True
    r.font.color.rgb = NAVY
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def add_steps(doc: Document, steps: list[str]) -> None:
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(step)
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, "1A2B4A")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)

    for i, row in enumerate(rows, 1):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            if i % 2 == 0:
                set_cell_shading(cell, "F5F7FA")
            p = cell.paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(10)

    doc.add_paragraph()


def build_manual() -> Document:
    doc = Document()
    configure_styles(doc)
    setup_page(doc.sections[0])

    add_cover(doc)

    # Sommaire
    doc.add_heading("Sommaire", level=1)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading("1. Objet et périmètre du document", level=1)
    doc.add_paragraph(
        "Le présent manuel décrit les procédures d'utilisation d'AvoLex, plateforme SaaS "
        "de gestion de cabinet d'avocats. Il s'adresse aux utilisateurs finaux : avocats, "
        "collaborateurs, secrétaires et administrateurs de cabinet."
    )
    doc.add_paragraph(
        "AvoLex couvre la gestion des clients, des dossiers, de l'agenda, des documents "
        "(GED), de la facturation et du tableau de bord. Chaque cabinet dispose d'un "
        "espace isolé (multi-tenant) : les données d'un cabinet ne sont jamais visibles "
        "par un autre."
    )

    add_info_box(
        doc,
        "Prérequis utilisateur",
        "Navigateur web récent (Chrome, Firefox, Edge). Compte utilisateur actif rattaché "
        "à au moins un cabinet. Connexion Internet stable.",
    )

    # 2. Accès
    doc.add_heading("2. Accès à l'application", level=1)

    doc.add_heading("2.1 Connexion", level=2)
    add_steps(
        doc,
        [
            "Ouvrir l'adresse de l'application (ex. https://votre-domaine.avolex.fr ou "
            "http://127.0.0.1:8000 en environnement local).",
            "Cliquer sur « Se connecter » ou accéder directement à /accounts/login/.",
            "Saisir votre adresse e-mail et votre mot de passe.",
            "Valider. Vous êtes redirigé vers le tableau de bord (/app/).",
        ],
    )

    doc.add_heading("2.2 Première inscription (création de cabinet)", level=2)
    add_steps(
        doc,
        [
            "Accéder à /accounts/register/.",
            "Renseigner vos informations personnelles et le nom du cabinet.",
            "Choisir un mot de passe robuste (8 caractères minimum, lettres et chiffres).",
            "Valider le formulaire. Vous devenez Propriétaire (owner) du cabinet créé.",
        ],
    )

    doc.add_heading("2.3 Mot de passe oublié", level=2)
    add_steps(
        doc,
        [
            "Sur la page de connexion, cliquer sur « Mot de passe oublié ».",
            "Saisir votre adresse e-mail et valider.",
            "Consulter votre boîte mail et suivre le lien reçu.",
            "Définir un nouveau mot de passe puis vous reconnecter.",
        ],
    )

    doc.add_heading("2.4 Accepter une invitation", level=2)
    add_steps(
        doc,
        [
            "Recevoir l'e-mail d'invitation envoyé par un membre autorisé du cabinet.",
            "Cliquer sur le lien contenant le jeton d'invitation "
            "(/accounts/invitations/<token>/).",
            "Si vous n'avez pas encore de compte : créer un mot de passe.",
            "Si vous avez déjà un compte : vous connecter pour accepter l'invitation.",
            "Vous accédez au cabinet avec le rôle qui vous a été attribué.",
        ],
    )

    doc.add_heading("2.5 Changer de cabinet actif", level=2)
    doc.add_paragraph(
        "Un utilisateur peut appartenir à plusieurs cabinets. Le sélecteur de cabinet "
        "dans l'interface permet de basculer le contexte actif. Toutes les listes et "
        "actions portent sur le cabinet sélectionné."
    )
    add_steps(
        doc,
        [
            "Ouvrir le menu ou le sélecteur de cabinet (barre latérale).",
            "Choisir le cabinet souhaité.",
            "Confirmer si nécessaire (action POST vers /cabinets/switch/).",
        ],
    )

    # 3. Rôles
    doc.add_heading("3. Rôles et permissions", level=1)
    doc.add_paragraph(
        "Chaque membre possède un rôle au sein du cabinet. Les permissions déterminent "
        "ce qu'il peut consulter, créer, modifier ou supprimer."
    )

    add_table(
        doc,
        ["Rôle", "Description", "Droits principaux"],
        [
            [
                "Propriétaire (owner)",
                "Administrateur du cabinet",
                "Tous les droits : gestion membres, facturation, cabinet, invitations",
            ],
            [
                "Avocat (lawyer)",
                "Avocat du cabinet",
                "CRUD métier, invitations de nouveaux membres",
            ],
            [
                "Collaborateur (associate)",
                "Avocat junior ou stagiaire",
                "Consultation, création et modification (pas de suppression)",
            ],
            [
                "Secrétaire (secretary)",
                "Assistant(e) administratif(ve)",
                "Consultation, création et modification (pas de suppression)",
            ],
            [
                "Lecture seule (read_only)",
                "Consultation uniquement",
                "Affichage des données sans modification",
            ],
        ],
    )

    add_table(
        doc,
        ["Permission", "Owner", "Avocat", "Collab.", "Secr.", "Lecture"],
        [
            ["Consulter (view)", "✓", "✓", "✓", "✓", "✓"],
            ["Créer (add)", "✓", "✓", "✓", "✓", "—"],
            ["Modifier (change)", "✓", "✓", "✓", "✓", "—"],
            ["Supprimer (delete)", "✓", "✓", "—", "—", "—"],
            ["Inviter (invite)", "✓", "✓", "—", "—", "—"],
            ["Gérer membres", "✓", "—", "—", "—", "—"],
            ["Gérer facturation", "✓", "—", "—", "—", "—"],
            ["Gérer cabinet", "✓", "—", "—", "—", "—"],
        ],
    )

    doc.add_heading("3.1 Inviter un membre", level=2)
    add_steps(
        doc,
        [
            "Accéder à /cabinets/invitations/ (menu Équipe ou Invitations).",
            "Cliquer sur « Inviter un membre ».",
            "Saisir l'adresse e-mail et sélectionner le rôle.",
            "Envoyer l'invitation. Le destinataire reçoit un lien par e-mail.",
            "Suivre le statut des invitations en attente depuis la même page.",
        ],
    )

    # 4. Navigation
    doc.add_heading("4. Navigation et tableau de bord", level=1)
    doc.add_paragraph(
        "Après connexion, le tableau de bord (/app/) affiche une vue synthétique de "
        "l'activité du cabinet."
    )

    add_table(
        doc,
        ["Indicateur", "Signification"],
        [
            ["Clients", "Nombre total de clients enregistrés"],
            ["Dossiers actifs", "Dossiers ouverts, en cours ou en attente"],
            ["Événements à venir", "Rendez-vous et échéances non terminés"],
            ["Heures non facturées", "Montant estimé des temps non encore facturés"],
            ["Factures impayées", "Factures envoyées ou en retard de paiement"],
            ["CA du mois", "Chiffre d'affaires facturé sur le mois en cours"],
        ],
    )

    add_table(
        doc,
        ["Module", "URL", "Description"],
        [
            ["Tableau de bord", "/app/", "Vue d'ensemble et KPIs"],
            ["Clients", "/clients/", "Fichier clients"],
            ["Dossiers", "/dossiers/", "Affaires et contentieux"],
            ["Agenda", "/agenda/", "Événements, échéances et tâches"],
            ["Documents", "/documents/", "GED par dossier"],
            ["Facturation", "/billing/", "Temps, débours, factures"],
            ["Invitations", "/cabinets/invitations/", "Gestion de l'équipe"],
        ],
    )

    # 5. Clients
    doc.add_heading("5. Gestion des clients", level=1)

    doc.add_heading("5.1 Consulter la liste des clients", level=2)
    add_steps(
        doc,
        [
            "Menu « Clients » ou URL /clients/.",
            "Utiliser la barre de recherche (nom, e-mail, téléphone, raison sociale).",
            "Parcourir les pages via la pagination en bas de liste.",
            "Cliquer sur un client pour ouvrir sa fiche détaillée.",
        ],
    )

    doc.add_heading("5.2 Créer un client", level=2)
    add_steps(
        doc,
        [
            "Depuis /clients/, cliquer sur « Nouveau client ».",
            "Choisir le type : Personne physique ou Société.",
            "Renseigner les panneaux Identité, Coordonnées, Adresse et Notes.",
            "Pour une personne : nom obligatoire. Pour une société : raison sociale obligatoire.",
            "Enregistrer. Le client est rattaché au cabinet actif.",
        ],
    )

    doc.add_heading("5.3 Modifier ou supprimer un client", level=2)
    add_steps(
        doc,
        [
            "Ouvrir la fiche client.",
            "Cliquer sur « Modifier » pour mettre à jour les informations.",
            "La suppression est réservée aux rôles Propriétaire et Avocat.",
            "Vérifier l'absence de dossiers liés avant toute suppression définitive.",
        ],
    )

    doc.add_heading("5.4 Fiche client", level=2)
    doc.add_paragraph(
        "La fiche présente l'identité du client, ses coordonnées, les notes internes "
        "et la liste des dossiers associés. Elle constitue le point d'entrée pour "
        "créer un nouveau dossier rattaché à ce client."
    )

    # 6. Dossiers
    doc.add_heading("6. Gestion des dossiers", level=1)

    doc.add_heading("6.1 Consulter et filtrer les dossiers", level=2)
    add_steps(
        doc,
        [
            "Accéder à /dossiers/.",
            "Filtrer par statut (ouvert, en cours, en attente, clos, archivé).",
            "Rechercher par titre, référence ou client.",
            "Ouvrir un dossier pour consulter le détail et l'historique des actions.",
        ],
    )

    doc.add_heading("6.2 Créer un dossier", level=2)
    add_steps(
        doc,
        [
            "Cliquer sur « Nouveau dossier ».",
            "Sélectionner le client concerné.",
            "Renseigner le titre, la matière, la juridiction et la partie adverse si applicable.",
            "Assigner l'avocat responsable et définir le statut initial.",
            "Enregistrer. Une référence unique est générée (format DOS-AAAA-NNNNN).",
        ],
    )

    doc.add_heading("6.3 Historique et suivi", level=2)
    doc.add_paragraph(
        "Chaque dossier conserve un journal des actions (MatterAction) : création, "
        "modification de statut, notes ajoutées. Consultez cet historique depuis la "
        "fiche dossier pour assurer la traçabilité du suivi."
    )

    doc.add_heading("6.4 Clôturer un dossier", level=2)
    add_steps(
        doc,
        [
            "Ouvrir le dossier et cliquer sur « Modifier ».",
            "Passer le statut à « Clos » ou « Archivé ».",
            "Renseigner la date de clôture si nécessaire.",
            "Vérifier que les temps et débours sont facturés avant clôture définitive.",
        ],
    )

    # 7. Agenda
    doc.add_heading("7. Agenda, événements et tâches", level=1)

    doc.add_heading("7.1 Types d'éléments", level=2)
    add_table(
        doc,
        ["Type", "Usage"],
        [
            ["Rendez-vous", "Audience, réunion client, conférence"],
            ["Échéance", "Date limite procédurale ou réglementaire"],
            ["Tâche", "Action à réaliser, cochable comme terminée"],
        ],
    )

    doc.add_heading("7.2 Créer un événement ou une tâche", level=2)
    add_steps(
        doc,
        [
            "Accéder à /agenda/ puis « Nouvel événement ».",
            "Choisir le type (rendez-vous, échéance, tâche).",
            "Renseigner titre, dates/heures, lieu et description.",
            "Associer optionnellement un dossier.",
            "Définir une date de rappel (remind_at) pour recevoir une notification.",
            "Enregistrer.",
        ],
    )

    doc.add_heading("7.3 Rappels automatiques", level=2)
    doc.add_paragraph(
        "Les rappels programmés sont traités par le service Celery (tâches planifiées). "
        "Assurez-vous que le worker Celery et Celery Beat sont actifs en production. "
        "Sans eux, les rappels ne seront pas envoyés."
    )

    doc.add_heading("7.4 Marquer une tâche comme terminée", level=2)
    add_steps(
        doc,
        [
            "Ouvrir l'événement de type « Tâche ».",
            "Cocher « Terminé » ou utiliser l'action rapide depuis la liste.",
            "La tâche disparaît des indicateurs « à venir » du tableau de bord.",
        ],
    )

    # 8. Documents
    doc.add_heading("8. Gestion documentaire (GED)", level=1)

    doc.add_heading("8.1 Principes", level=2)
    doc.add_paragraph(
        "Les documents sont stockés dans un espace privé (private_media/), hors "
        "accès web direct. Chaque document est rattaché à un dossier et versionné."
    )

    doc.add_heading("8.2 Déposer un document", level=2)
    add_steps(
        doc,
        [
            "Accéder à /documents/ puis « Déposer un document ».",
            "Sélectionner le dossier cible.",
            "Choisir le fichier (respecter la taille maximale autorisée).",
            "Renseigner titre, description et tags (séparés par des virgules).",
            "Valider. La version 1 est créée automatiquement.",
        ],
    )

    doc.add_heading("8.3 Versions et métadonnées", level=2)
    add_steps(
        doc,
        [
            "Ouvrir la fiche document depuis /documents/.",
            "Consulter l'historique des versions.",
            "Ajouter une nouvelle version si le fichier a évolué.",
            "Modifier les métadonnées (titre, description, tags) sans changer le fichier.",
            "Télécharger ou prévisualiser une version spécifique.",
        ],
    )

    doc.add_heading("8.4 Suppression", level=2)
    doc.add_paragraph(
        "La suppression d'un document est réservée aux rôles autorisés (Propriétaire, "
        "Avocat). Elle est définitive pour l'ensemble des versions."
    )

    # 9. Facturation
    doc.add_heading("9. Facturation", level=1)

    doc.add_heading("9.1 Vue d'ensemble", level=2)
    doc.add_paragraph(
        "Le module Facturation (/billing/) regroupe les temps passés, les débours, "
        "les factures et le timer de saisie. Seul le Propriétaire peut gérer les "
        "paramètres avancés de facturation du cabinet."
    )

    doc.add_heading("9.2 Saisir du temps", level=2)
    add_steps(
        doc,
        [
            "Aller dans Facturation → Temps passés → « Saisir du temps ».",
            "Sélectionner le dossier et décrire la prestation.",
            "Indiquer la durée en minutes et le taux horaire.",
            "Cocher « Facturable » si le temps doit apparaître sur une facture.",
            "Enregistrer.",
        ],
    )

    doc.add_heading("9.3 Utiliser le timer", level=2)
    add_steps(
        doc,
        [
            "Depuis le hub Facturation, démarrer un timer sur un dossier.",
            "Travailler sur le dossier concerné.",
            "Arrêter le timer : une entrée de temps est créée automatiquement.",
            "Vérifier et ajuster la description ou la durée si nécessaire.",
        ],
    )

    doc.add_heading("9.4 Enregistrer un débours", level=2)
    add_steps(
        doc,
        [
            "Facturation → Débours → « Nouveau débours ».",
            "Sélectionner le dossier, saisir description et montant.",
            "Indiquer si le débours est refacturable au client.",
            "Enregistrer.",
        ],
    )

    doc.add_heading("9.5 Émettre une facture", level=2)
    add_steps(
        doc,
        [
            "Facturation → Factures → « Nouvelle facture ».",
            "Choisir le client et le dossier.",
            "Sélectionner les temps et débours non encore facturés.",
            "Vérifier le montant HT, la TVA et le total TTC.",
            "Générer la facture (PDF via WeasyPrint en production).",
            "Passer le statut à « Envoyée » après envoi au client.",
        ],
    )

    add_table(
        doc,
        ["Statut facture", "Signification"],
        [
            ["Brouillon", "En cours de préparation, modifiable"],
            ["Envoyée", "Transmise au client, en attente de paiement"],
            ["Payée", "Règlement reçu et enregistré"],
            ["En retard", "Échéance dépassée sans paiement"],
            ["Annulée", "Facture annulée, non comptabilisée"],
        ],
    )

    # 10. Sécurité
    doc.add_heading("10. Sécurité, confidentialité et bonnes pratiques", level=1)

    doc.add_heading("10.1 Isolation multi-tenant", level=2)
    doc.add_paragraph(
        "AvoLex isole strictement les données par cabinet. Le middleware CabinetMiddleware "
        "filtre toutes les requêtes : sans cabinet actif valide, aucune donnée métier "
        "n'est accessible (principe fail-closed)."
    )

    doc.add_heading("10.2 Mots de passe et sessions", level=2)
    add_steps(
        doc,
        [
            "Utiliser un mot de passe unique et complexe pour AvoLex.",
            "Ne jamais partager vos identifiants.",
            "Se déconnecter sur les postes partagés (/accounts/logout/).",
            "Renouveler le mot de passe en cas de suspicion de compromission.",
        ],
    )

    doc.add_heading("10.3 Documents et RGPD", level=2)
    doc.add_paragraph(
        "Les documents clients sont des données sensibles. Ne les téléchargez que sur "
        "des postes sécurisés. Respectez la durée de conservation définie par votre "
        "cabinet. En cas de demande d'effacement, coordonnez-vous avec le Propriétaire "
        "avant toute suppression."
    )

    doc.add_heading("10.4 Rôles minimum", level=2)
    doc.add_paragraph(
        "Attribuez le rôle le plus restrictif compatible avec les missions : un stagiaire "
        "n'a pas besoin des droits de suppression ; un expert-comptable externe peut "
        "recevoir un accès lecture seule."
    )

    # 11. API
    doc.add_heading("11. API REST (usage avancé)", level=1)
    doc.add_paragraph(
        "Une API REST (Django REST Framework) est disponible sous /api/v1/ pour "
        "l'intégration avec des outils tiers."
    )
    add_table(
        doc,
        ["Endpoint", "Ressource"],
        [
            ["/api/v1/clients/", "Clients du cabinet actif"],
            ["/api/v1/matters/", "Dossiers"],
            ["/api/v1/events/", "Événements agenda"],
        ],
    )
    doc.add_paragraph(
        "L'authentification repose sur la session Django (cookie après connexion web). "
        "Les mêmes règles de permissions s'appliquent que dans l'interface."
    )

    # 12. Dépannage
    doc.add_heading("12. Dépannage courant", level=1)
    add_table(
        doc,
        ["Problème", "Action recommandée"],
        [
            ["Impossible de se connecter", "Vérifier e-mail/mot de passe ; utiliser la réinitialisation"],
            ["Module inaccessible", "Vérifier votre rôle ; contacter le Propriétaire"],
            ["Données absentes", "Vérifier le cabinet actif dans le sélecteur"],
            ["Rappels non reçus", "Contacter l'administrateur (Celery worker/beat)"],
            ["PDF facture indisponible", "Vérifier WeasyPrint côté serveur"],
        ],
    )

    # 13. Compte démo
    doc.add_heading("13. Environnement de démonstration", level=1)
    doc.add_paragraph(
        "Pour la formation ou les tests, un jeu de données de démonstration peut être "
        "chargé par l'administrateur technique :"
    )
    add_table(
        doc,
        ["Élément", "Valeur"],
        [
            ["Commande", "python manage.py seed_demo"],
            ["E-mail", "ngamika@gmail.com"],
            ["Mot de passe", "@12345678"],
            ["URL application", "http://127.0.0.1:8000/app/"],
        ],
    )

    # 14. Contacts
    doc.add_heading("14. Support et évolution", level=1)
    doc.add_paragraph(
        "Pour toute demande d'évolution fonctionnelle, de correction ou de formation "
        "complémentaire, contactez l'administrateur AvoLex de votre cabinet ou "
        "l'équipe technique responsable du déploiement."
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— Fin du document —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = GRAY

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_manual()
    doc.save(str(OUTPUT))
    print(f"Manuel généré : {OUTPUT}")


if __name__ == "__main__":
    main()
